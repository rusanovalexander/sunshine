"""
=================================================
Stage 1: Enhanced Document Preprocessing (v2)
=================================================
Improvements:
- Better structure preservation (sections, headers)
- Table detection and specialized handling
- Semantic chunking with overlap
- Metadata extraction
"""

import os
import gc
import zipfile
import shutil
import tempfile
import torch
import pandas as pd
import logging
import re
import json
import hashlib
import argparse
from io import BytesIO
from PIL import Image
from typing import List, Dict, Tuple, Optional

# Compression
try:
    import py7zr
except ImportError:
    py7zr = None

# PDF & OCR
import pymupdf as fitz
import easyocr
from langdetect import detect

# Office Formats
import extract_msg
from docx import Document
from bs4 import BeautifulSoup

# Excel
try:
    import openpyxl
except ImportError:
    openpyxl = None

# VLM
try:
    from transformers import AutoProcessor, Qwen2_5_VLForConditionalGeneration, BitsAndBytesConfig
except ImportError:
    Qwen2_5_VLForConditionalGeneration = None
    BitsAndBytesConfig = None

# GPU Optimization
try:
    from .gpu_optimizer import (
        load_vlm_optimized, log_gpu_memory, aggressive_memory_cleanup,
        get_gpu_memory_stats, unload_model
    )
    GPU_OPTIMIZER_AVAILABLE = True
except ImportError:
    GPU_OPTIMIZER_AVAILABLE = False

from .config import (
    VLM_MODEL_PATH, ARCHIVE_PATH, EXTRACTED_SOURCE_DIR, 
    PREPROCESSED_DATA_DIR, CHUNKS_DIR, CHUNK_SIZE, CHUNK_OVERLAP, DPI
)

# =====================================================================
# LOGGING
# =====================================================================
logging.basicConfig(
    level=logging.INFO, 
    format='%(asctime)s - Preprocess - %(levelname)s - %(message)s',
    datefmt='%H:%M:%S'
)
logger = logging.getLogger(__name__)

IMAGE_OCR_CACHE = {}

# =====================================================================
# STRUCTURE-PRESERVING TEXT EXTRACTION
# =====================================================================

class DocumentSection:
    """Represents a section of a document with metadata."""
    def __init__(self, content: str, section_type: str = "text", 
                 page_num: int = None, header: str = None):
        self.content = content
        self.section_type = section_type  # "text", "table", "header", "list"
        self.page_num = page_num
        self.header = header
    
    def to_dict(self):
        return {
            "content": self.content,
            "type": self.section_type,
            "page": self.page_num,
            "header": self.header
        }


def detect_section_headers(text: str) -> List[Tuple[int, str]]:
    """Detect potential section headers in text."""
    headers = []
    lines = text.split('\n')
    
    header_patterns = [
        r'^(?:SECTION|ARTICLE|PART|SCHEDULE|ANNEX|APPENDIX|EXHIBIT)\s+[\dIVXA-Z]+',
        r'^[\d]+\.[\d]*\s+[A-Z][A-Za-z\s]+$',
        r'^[A-Z][A-Z\s]{10,}$',  # ALL CAPS lines
        r'^\d+\.\s+[A-Z]',  # Numbered sections
    ]
    
    for i, line in enumerate(lines):
        stripped = line.strip()
        if len(stripped) < 3 or len(stripped) > 100:
            continue
        for pattern in header_patterns:
            if re.match(pattern, stripped):
                headers.append((i, stripped))
                break
    
    return headers


def detect_tables_in_text(text: str) -> List[Tuple[int, int, str]]:
    """Detect table-like structures in text and return their positions."""
    tables = []
    lines = text.split('\n')
    
    in_table = False
    table_start = 0
    table_lines = []
    
    for i, line in enumerate(lines):
        # Detect table indicators
        is_table_line = (
            '|' in line or 
            line.count('\t') >= 2 or
            re.match(r'^\s*[\d,\.]+\s+[\d,\.]+\s+[\d,\.]+', line) or  # Number columns
            re.match(r'^\s*[-=+]+\s*$', line)  # Separator lines
        )
        
        if is_table_line and not in_table:
            in_table = True
            table_start = i
            table_lines = [line]
        elif is_table_line and in_table:
            table_lines.append(line)
        elif not is_table_line and in_table:
            if len(table_lines) >= 3:  # Minimum table size
                tables.append((table_start, i, '\n'.join(table_lines)))
            in_table = False
            table_lines = []
    
    return tables


def structure_document(raw_text: str, source_file: str) -> Dict:
    """Parse document into structured sections."""
    sections = []
    
    # Detect headers
    headers = detect_section_headers(raw_text)
    header_positions = {pos: header for pos, header in headers}
    
    # Detect tables
    tables = detect_tables_in_text(raw_text)
    table_ranges = [(start, end) for start, end, _ in tables]
    
    lines = raw_text.split('\n')
    current_header = None
    current_content = []
    current_start = 0
    
    def is_in_table(line_num):
        return any(start <= line_num < end for start, end in table_ranges)
    
    for i, line in enumerate(lines):
        # Check if this is a header
        if i in header_positions:
            # Save previous section
            if current_content:
                sections.append(DocumentSection(
                    content='\n'.join(current_content),
                    section_type="text",
                    header=current_header
                ))
            current_header = header_positions[i]
            current_content = []
            current_start = i
        elif not is_in_table(i):
            current_content.append(line)
    
    # Add final section
    if current_content:
        sections.append(DocumentSection(
            content='\n'.join(current_content),
            section_type="text",
            header=current_header
        ))
    
    # Add tables as separate sections
    for start, end, table_content in tables:
        # Find which header this table belongs to
        table_header = None
        for pos, header in sorted(headers, reverse=True):
            if pos < start:
                table_header = header
                break
        
        sections.append(DocumentSection(
            content=table_content,
            section_type="table",
            header=table_header
        ))
    
    return {
        "source_file": source_file,
        "total_sections": len(sections),
        "sections": [s.to_dict() for s in sections],
        "headers_found": [h for _, h in headers],
        "tables_found": len(tables)
    }


# =====================================================================
# SEMANTIC CHUNKING
# =====================================================================

def chunk_text_with_overlap(text: str, tokenizer, chunk_size: int = CHUNK_SIZE, 
                            overlap: int = CHUNK_OVERLAP) -> List[Dict]:
    """Create overlapping chunks that respect sentence boundaries."""
    chunks = []
    
    # Split into sentences (approximate)
    sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', text)
    
    current_chunk = []
    current_tokens = 0
    chunk_id = 0
    
    for sentence in sentences:
        sentence_tokens = len(tokenizer.encode(sentence, add_special_tokens=False))
        
        if current_tokens + sentence_tokens > chunk_size and current_chunk:
            # Save current chunk
            chunk_text = ' '.join(current_chunk)
            chunks.append({
                "id": chunk_id,
                "text": chunk_text,
                "token_count": current_tokens,
                "sentence_count": len(current_chunk)
            })
            chunk_id += 1
            
            # Start new chunk with overlap
            # Keep last few sentences for context
            overlap_sentences = []
            overlap_tokens = 0
            for s in reversed(current_chunk):
                s_tokens = len(tokenizer.encode(s, add_special_tokens=False))
                if overlap_tokens + s_tokens <= overlap:
                    overlap_sentences.insert(0, s)
                    overlap_tokens += s_tokens
                else:
                    break
            
            current_chunk = overlap_sentences + [sentence]
            current_tokens = overlap_tokens + sentence_tokens
        else:
            current_chunk.append(sentence)
            current_tokens += sentence_tokens
    
    # Don't forget the last chunk
    if current_chunk:
        chunks.append({
            "id": chunk_id,
            "text": ' '.join(current_chunk),
            "token_count": current_tokens,
            "sentence_count": len(current_chunk)
        })
    
    return chunks


# =====================================================================
# VLM / OCR
# =====================================================================

def initialize_vlm(model_path: str, use_flash_attention: bool = False):
    """Load Qwen2.5-VL for OCR with GPU optimization."""
    logger.info(f"Loading VLM: {model_path}")
    
    # Use optimized loader if available
    if GPU_OPTIMIZER_AVAILABLE:
        try:
            model, processor = load_vlm_optimized(
                model_path,
                use_flash_attention=use_flash_attention,
                max_memory_gb=6.0  # VLM is smaller, leave room for LLM later
            )
            return model, processor
        except Exception as e:
            logger.error(f"Optimized VLM load failed: {e}")
    
    # Fallback to standard loading
    if not Qwen2_5_VLForConditionalGeneration:
        logger.error("Transformers library not available")
        return None, None
    
    bnb_config = BitsAndBytesConfig(
        load_in_4bit=True,
        bnb_4bit_compute_dtype=torch.float16,
        bnb_4bit_quant_type="nf4"
    )
    
    attn_impl = "flash_attention_2" if use_flash_attention else "sdpa"
    
    try:
        model = Qwen2_5_VLForConditionalGeneration.from_pretrained(
            model_path,
            quantization_config=bnb_config,
            device_map="auto",
            attn_implementation=attn_impl,
            trust_remote_code=True
        )
        processor = AutoProcessor.from_pretrained(model_path, trust_remote_code=True)
        return model, processor
    except Exception as e:
        logger.error(f"Failed to load VLM: {e}")
        return None, None


def run_vlm_ocr(model, processor, image: Image.Image, prompt: str = "Extract all text from this image.") -> str:
    """Run VLM inference for OCR."""
    try:
        messages = [
            {"role": "user", "content": [
                {"type": "image", "image": image},
                {"type": "text", "text": prompt}
            ]}
        ]
        
        text_input = processor.apply_chat_template(
            messages, tokenize=False, add_generation_prompt=True, enable_thinking=False
        )
        inputs = processor(
            text=[text_input],
            images=[image],
            padding=True,
            return_tensors="pt"
        ).to(model.device)
        
        with torch.no_grad():
            output_ids = model.generate(**inputs, max_new_tokens=2048, do_sample=False)
        
        generated_ids = output_ids[:, inputs.input_ids.shape[1]:]
        return processor.batch_decode(generated_ids, skip_special_tokens=True)[0]
    except Exception as e:
        logger.error(f"VLM OCR error: {e}")
        return ""


def ocr_image(image: Image.Image, ocr_engine, context: str = "") -> str:
    """Route image to appropriate OCR engine."""
    if not ocr_engine:
        return ""
    
    try:
        if isinstance(ocr_engine, tuple):  # VLM
            prompt = f"Extract all text from this {context}. Preserve formatting and structure."
            return run_vlm_ocr(ocr_engine[0], ocr_engine[1], image, prompt)
        elif isinstance(ocr_engine, easyocr.Reader):
            import numpy as np
            result = ocr_engine.readtext(np.array(image))
            return "\n".join([r[1] for r in result])
    except Exception as e:
        logger.error(f"OCR error: {e}")
    return ""


# =====================================================================
# FILE EXTRACTORS
# =====================================================================

def clean_text(text: str) -> str:
    """Clean extracted text while preserving structure."""
    if not text:
        return ""
    
    # Remove pandas artifacts
    text = re.sub(r'\b(NaN|Unnamed:\s*\d+)\b', '', text)
    
    # Normalize whitespace within lines
    text = re.sub(r'[ \t]+', ' ', text)
    
    # Limit consecutive newlines
    text = re.sub(r'\n{4,}', '\n\n\n', text)
    
    # Remove lines that are just whitespace or very short noise
    lines = []
    for line in text.split('\n'):
        stripped = line.strip()
        if stripped and (len(stripped) > 2 or stripped in ['|', '-', '+']):
            lines.append(line)
    
    return '\n'.join(lines).strip()


def extract_from_pdf(pdf_path: str, ocr_engine) -> Tuple[str, Dict]:
    """Extract text from PDF with page markers and OCR fallback."""
    full_text = []
    metadata = {"pages": 0, "ocr_pages": [], "has_images": False}
    
    try:
        doc = fitz.open(pdf_path)
        metadata["pages"] = len(doc)
        
        for page_num, page in enumerate(doc):
            page_text = page.get_text().strip()
            
            # Check if page needs OCR
            if len(page_text) < 100 and ocr_engine:
                pix = page.get_pixmap(dpi=DPI)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                ocr_text = ocr_image(img, ocr_engine, f"PDF page {page_num + 1}")
                
                if ocr_text.strip():
                    page_text = ocr_text
                    metadata["ocr_pages"].append(page_num + 1)
            
            full_text.append(f"\n{'='*60}\n[PAGE {page_num + 1}]\n{'='*60}\n{page_text}")
        
        return '\n'.join(full_text), metadata
    except Exception as e:
        logger.error(f"PDF extraction error {pdf_path}: {e}")
        return "", metadata


def extract_from_excel(xls_path: str, ocr_engine) -> Tuple[str, Dict]:
    """Extract Excel with tables as markdown and embedded images."""
    parts = []
    metadata = {"sheets": [], "images_processed": 0}
    
    try:
        xls = pd.ExcelFile(xls_path)
        
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name, header=None)
            df = df.dropna(how='all').fillna('')
            
            if not df.empty:
                metadata["sheets"].append(sheet_name)
                
                # Convert to markdown table
                md_table = df.to_markdown(index=False)
                parts.append(f"\n[SHEET: {sheet_name}]\n{md_table}")
        
        # Process embedded images
        if xls_path.endswith('.xlsx') and ocr_engine and openpyxl:
            wb = openpyxl.load_workbook(xls_path, data_only=True)
            
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                if hasattr(ws, '_images'):
                    for i, img_container in enumerate(ws._images):
                        try:
                            img_bytes = img_container._data()
                            img_hash = hashlib.md5(img_bytes).hexdigest()
                            
                            if img_hash in IMAGE_OCR_CACHE:
                                if IMAGE_OCR_CACHE[img_hash]:
                                    parts.append(f"[IMAGE in {sheet_name}]\n{IMAGE_OCR_CACHE[img_hash]}")
                                continue
                            
                            img = Image.open(BytesIO(img_bytes)).convert("RGB")
                            w, h = img.size
                            
                            # Skip tiny/decorative images
                            if w < 50 or h < 50 or max(w, h) / min(w, h) > 20:
                                IMAGE_OCR_CACHE[img_hash] = ""
                                continue
                            
                            text = ocr_image(img, ocr_engine, "embedded Excel image")
                            IMAGE_OCR_CACHE[img_hash] = text
                            metadata["images_processed"] += 1
                            
                            if text.strip():
                                parts.append(f"[IMAGE in {sheet_name}]\n{text}")
                        except Exception as e:
                            logger.warning(f"Excel image error: {e}")
        
        return '\n'.join(parts), metadata
    except Exception as e:
        logger.error(f"Excel extraction error {xls_path}: {e}")
        return "", metadata


def extract_from_docx(docx_path: str) -> Tuple[str, Dict]:
    """Extract from Word document preserving structure."""
    parts = []
    metadata = {"paragraphs": 0, "tables": 0}
    
    try:
        doc = Document(docx_path)
        
        for para in doc.paragraphs:
            if para.text.strip():
                # Check if it's a heading
                if para.style and 'Heading' in para.style.name:
                    parts.append(f"\n## {para.text}\n")
                else:
                    parts.append(para.text)
                metadata["paragraphs"] += 1
        
        # Extract tables
        for table in doc.tables:
            metadata["tables"] += 1
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(' | '.join(row_data))
            parts.append(f"\n[TABLE]\n" + '\n'.join(table_data) + "\n")
        
        return '\n'.join(parts), metadata
    except Exception as e:
        logger.error(f"DOCX extraction error {docx_path}: {e}")
        return "", metadata


def extract_from_msg(msg_path: str) -> Tuple[str, Dict]:
    """Extract from Outlook MSG file."""
    metadata = {"subject": "", "sender": "", "has_attachments": False}
    
    try:
        msg = extract_msg.Message(msg_path)
        metadata["subject"] = msg.subject or ""
        metadata["sender"] = msg.sender or ""
        metadata["has_attachments"] = bool(msg.attachments)
        
        content = f"""[EMAIL]
Subject: {msg.subject}
From: {msg.sender}
Date: {msg.date}
---
{msg.body}
"""
        return content, metadata
    except Exception as e:
        logger.error(f"MSG extraction error {msg_path}: {e}")
        return "", metadata


def extract_file(file_path: str, ocr_engine) -> Tuple[str, Dict]:
    """Route file to appropriate extractor."""
    ext = os.path.splitext(file_path)[1].lower()
    
    extractors = {
        '.pdf': lambda: extract_from_pdf(file_path, ocr_engine),
        '.xlsx': lambda: extract_from_excel(file_path, ocr_engine),
        '.xls': lambda: extract_from_excel(file_path, ocr_engine),
        '.xlsm': lambda: extract_from_excel(file_path, ocr_engine),
        '.docx': lambda: extract_from_docx(file_path),
        '.msg': lambda: extract_from_msg(file_path),
    }
    
    if ext in extractors:
        return extractors[ext]()
    elif ext in ['.png', '.jpg', '.jpeg', '.bmp', '.tiff']:
        text = ocr_image(Image.open(file_path).convert("RGB"), ocr_engine, "image file")
        return text, {"type": "image"}
    elif ext == '.txt':
        with open(file_path, 'r', errors='ignore') as f:
            return f.read(), {"type": "text"}
    else:
        return "", {"type": "unsupported"}


# =====================================================================
# MAIN PROCESSING
# =====================================================================

def process_company(company_name: str, source_root: str, dest_root: str, 
                    chunks_root: str, ocr_engine, tokenizer) -> List[Dict]:
    """Process all files for a company."""
    source_dir = os.path.join(source_root, company_name)
    dest_dir = os.path.join(dest_root, company_name)
    chunks_dir = os.path.join(chunks_root, company_name)
    
    os.makedirs(dest_dir, exist_ok=True)
    os.makedirs(chunks_dir, exist_ok=True)
    
    manifest_entries = []
    
    for root, _, files in os.walk(source_dir):
        for filename in files:
            src_path = os.path.join(root, filename)
            rel_path = os.path.relpath(src_path, source_dir)
            safe_name = re.sub(r'[^a-zA-Z0-9_.-]', '_', rel_path)
            
            # Output paths
            text_file = os.path.join(dest_dir, f"{safe_name}.txt")
            struct_file = os.path.join(dest_dir, f"{safe_name}.structure.json")
            chunks_file = os.path.join(chunks_dir, f"{safe_name}.chunks.json")
            
            # Skip if already processed
            if os.path.exists(text_file) and os.path.exists(chunks_file):
                logger.info(f"Skipping (exists): {rel_path}")
                continue
            
            logger.info(f"Processing: {rel_path}")
            
            # Extract
            text, metadata = extract_file(src_path, ocr_engine)
            text = clean_text(text)
            
            if not text.strip():
                logger.warning(f"No text extracted: {rel_path}")
                continue
            
            # Save raw text
            with open(text_file, 'w', encoding='utf-8') as f:
                f.write(text)
            
            # Structure analysis
            structure = structure_document(text, rel_path)
            structure["extraction_metadata"] = metadata
            with open(struct_file, 'w', encoding='utf-8') as f:
                json.dump(structure, f, indent=2)
            
            # Create chunks
            if tokenizer:
                chunks = chunk_text_with_overlap(text, tokenizer)
                with open(chunks_file, 'w', encoding='utf-8') as f:
                    json.dump(chunks, f, indent=2)
            
            # Detect languages
            try:
                from langdetect import detect
                langs = [detect(text[:5000])]
            except:
                langs = ["unknown"]
            
            manifest_entries.append({
                "company": company_name,
                "original_file": rel_path,
                "text_file": os.path.join(company_name, f"{safe_name}.txt"),
                "structure_file": os.path.join(company_name, f"{safe_name}.structure.json"),
                "chunks_file": os.path.join(company_name, f"{safe_name}.chunks.json"),
                "languages": langs,
                "metadata": metadata
            })
    
    return manifest_entries


def main():
    parser = argparse.ArgumentParser(description="Document Preprocessing v2")
    parser.add_argument("--ocr_method", choices=['vlm', 'easyocr', 'none'], default='vlm')
    parser.add_argument("--flash_attention", action="store_true")
    parser.add_argument("--company", help="Process specific company only")
    args = parser.parse_args()
    
    # Extract archive if needed
    sentinel = os.path.join(EXTRACTED_SOURCE_DIR, ".ready")
    if not os.path.exists(sentinel):
        if os.path.exists(ARCHIVE_PATH):
            logger.info(f"Extracting archive: {ARCHIVE_PATH}")
            if ARCHIVE_PATH.endswith('.7z') and py7zr:
                with py7zr.SevenZipFile(ARCHIVE_PATH, 'r') as z:
                    z.extractall(EXTRACTED_SOURCE_DIR)
            elif ARCHIVE_PATH.endswith('.zip'):
                with zipfile.ZipFile(ARCHIVE_PATH, 'r') as z:
                    z.extractall(EXTRACTED_SOURCE_DIR)
            with open(sentinel, 'w') as f:
                f.write("done")
        elif not os.path.exists(EXTRACTED_SOURCE_DIR):
            logger.error("Source data not found")
            return
    
    # Initialize OCR
    ocr_engine = None
    if args.ocr_method == 'vlm':
        ocr_engine = initialize_vlm(VLM_MODEL_PATH, args.flash_attention)
    elif args.ocr_method == 'easyocr':
        try:
            # Try to load EasyOCR (requires pre-downloaded models for offline use)
            # Using multiple Latin languages for European financial documents
            ocr_engine = easyocr.Reader(
                ['en', 'fr', 'de', 'es', 'it', 'pt', 'nl'],  # Major EU languages
                download_enabled=False,
                gpu=True
            )
        except Exception as e:
            logger.warning(f"EasyOCR failed (offline server?): {e}")
            logger.warning("Falling back to VLM for OCR...")
            ocr_engine = initialize_vlm(VLM_MODEL_PATH, args.flash_attention)
            if ocr_engine is None:
                logger.warning("No OCR available - scanned documents won't be processed")
    
    # Initialize tokenizer for chunking
    from transformers import AutoTokenizer
    try:
        from .config import MODEL_PATH
        tokenizer = AutoTokenizer.from_pretrained(MODEL_PATH, trust_remote_code=True)
    except:
        tokenizer = None
        logger.warning("Could not load tokenizer for chunking")
    
    # Create output directories
    os.makedirs(PREPROCESSED_DATA_DIR, exist_ok=True)
    os.makedirs(CHUNKS_DIR, exist_ok=True)
    
    # Load or create manifest
    manifest_path = os.path.join(PREPROCESSED_DATA_DIR, "manifest.json")
    if os.path.exists(manifest_path):
        with open(manifest_path, 'r') as f:
            manifest = json.load(f)
    else:
        manifest = []
    
    # Get companies to process
    if args.company:
        companies = [args.company]
    else:
        companies = [d for d in os.listdir(EXTRACTED_SOURCE_DIR) 
                    if os.path.isdir(os.path.join(EXTRACTED_SOURCE_DIR, d)) and not d.startswith('.')]
    
    # Process
    for company in companies:
        logger.info(f"=== Company: {company} ===")
        new_entries = process_company(
            company, EXTRACTED_SOURCE_DIR, PREPROCESSED_DATA_DIR,
            CHUNKS_DIR, ocr_engine, tokenizer
        )
        
        if new_entries:
            manifest.extend(new_entries)
            with open(manifest_path, 'w') as f:
                json.dump(manifest, f, indent=2)
    
    # Unload VLM to free GPU memory for extraction stage
    if ocr_engine and isinstance(ocr_engine, tuple):
        logger.info("Unloading VLM model...")
        if GPU_OPTIMIZER_AVAILABLE:
            unload_model(ocr_engine[0])
            log_gpu_memory("After VLM unload: ")
        else:
            del ocr_engine
            gc.collect()
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
    
    logger.info(f"Preprocessing complete. {len(manifest)} files in manifest.")


if __name__ == "__main__":
    main()
